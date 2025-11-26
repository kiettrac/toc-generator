import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# --- CÁC HÀM XỬ LÝ WORD (CORE) ---
def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_toc_field(paragraph):
    run = paragraph.add_run()
    
    fldChar_begin = create_element('w:fldChar')
    create_attribute(fldChar_begin, 'w:fldCharType', 'begin')
    run._r.append(fldChar_begin)
    
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instrText)
    
    fldChar_sep = create_element('w:fldChar')
    create_attribute(fldChar_sep, 'w:fldCharType', 'separate')
    run._r.append(fldChar_sep)
    
    fldChar_end = create_element('w:fldChar')
    create_attribute(fldChar_end, 'w:fldCharType', 'end')
    run._r.append(fldChar_end)

def add_page_number_to_footer(section, position='center'):
    footer = section.footer
    footer.is_linked_to_previous = False
    
    if len(footer.paragraphs) == 0:
        paragraph = footer.add_paragraph()
    else:
        paragraph = footer.paragraphs[0]
    
    if position == 'left':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif position == 'right':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    paragraph.clear()
    run = paragraph.add_run()
    
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')
    
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

# --- GIAO DIỆN NGƯỜI DÙNG ---
st.set_page_config(page_title="Chèn Mục Lục Tùy Chọn", layout="wide")

# Custom CSS cho footer
st.markdown("""
<style>
.footer {
    position: fixed;
    bottom: 0;
    right: 20px;
    padding: 10px;
    color: #666;
    font-size: 18px;
    font-weight: bold;
    font-style: italic;
    text-align: center;
}
</style>
<div class="footer">
    Develop By TracTuanKiet
</div>
""", unsafe_allow_html=True)

st.title("📄 Tool Chèn Mục Lục Vào Vị Trí Bất Kỳ")
st.write("Upload file, xem nội dung, và chọn vị trí muốn chèn Mục lục.")

uploaded_file = st.file_uploader("Chọn file Word (.docx)", type=["docx"])

if uploaded_file is not None:
    doc = Document(uploaded_file)
    
    st.divider()
    st.subheader("👀 Xem trước nội dung & Chọn vị trí chèn")
    
    # Tạo danh sách các đoạn văn
    preview_options = []
    paragraph_map = {}
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if len(text) > 0:
            label = f"Đoạn {i}: {text[:80]}..."
            preview_options.append(label)
            paragraph_map[label] = i
    
    # Giao diện chia 2 cột
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.info("Danh sách các đoạn văn tìm thấy trong file:")
        full_text_preview = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        st.text_area("Nội dung file (Review)", full_text_preview, height=300)
    
    with col2:
        st.warning("🎯 Chọn đoạn văn bạn muốn chèn Mục Lục vào TRƯỚC nó:")
        selected_label = st.selectbox("Chọn vị trí:", options=preview_options)
        target_index = paragraph_map.get(selected_label)
        
        # Thêm tùy chọn vị trí số trang
        st.divider()
        st.write("📍 **Chọn vị trí hiển thị số trang:**")
        
        col_pos1, col_pos2, col_pos3 = st.columns(3)
        
        with col_pos1:
            page_position = st.radio(
                "Vị trí số trang",
                options=['left', 'center', 'right'],
                format_func=lambda x: {'left': '← Trái', 'center': '⊙ Giữa', 'right': '→ Phải'}[x],
                index=1  # Mặc định là center
            )
    
    if st.button("🚀 Chèn Mục Lục & Tạo File", type="primary"):
        try:
            # Lấy đoạn văn mục tiêu
            target_paragraph = doc.paragraphs[target_index]
            
            # Chèn tiêu đề "MỤC LỤC"
            p_title = target_paragraph.insert_paragraph_before("MỤC LỤC")
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_title.runs[0].bold = True
            p_title.runs[0].font.size = 180000
            
            # Chèn Code TOC
            p_toc = target_paragraph.insert_paragraph_before("")
            add_toc_field(p_toc)
            
            # Chèn Ngắt trang
            p_break = target_paragraph.insert_paragraph_before("")
            p_break_element = p_break._p
            sectPr = create_element('w:sectPr')
            type_element = create_element('w:type')
            create_attribute(type_element, 'w:val', 'nextPage')
            sectPr.append(type_element)
            p_break_element.get_or_add_pPr().append(sectPr)
            
            # Lưu tạm và Load lại
            buffer_temp = BytesIO()
            doc.save(buffer_temp)
            buffer_temp.seek(0)
            doc_v2 = Document(buffer_temp)
            
            # Thêm số trang với vị trí đã chọn
            if len(doc_v2.sections) > 1:
                content_section = doc_v2.sections[-1]
                
                # Reset trang về 1
                sectPr = content_section._sectPr
                pgNumType = create_element('w:pgNumType')
                create_attribute(pgNumType, 'w:start', '1')
                sectPr.append(pgNumType)
                
                # Thêm số trang với vị trí đã chọn
                add_page_number_to_footer(content_section, position=page_position)
            
            # Xuất file
            buffer_final = BytesIO()
            doc_v2.save(buffer_final)
            buffer_final.seek(0)
            
            st.success("✅ Xử lý xong! Tải file ở dưới:")
            st.download_button(
                label="⬇️ Tải file kết quả",
                data=buffer_final,
                file_name="File_Co_Muc_Luc_Va_So_Trang.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.info("🔔 Nhắc lại: Mở file → Ctrl+A → F9 → Update entire table để hiện số liệu.")
            
        except Exception as e:

            st.error(f"Có lỗi xảy ra: {e}")
